from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from bs4 import BeautifulSoup
import pandas as pd
import pandas as pd
import streamlit as st
from docx import Document
import os
import time
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.common.by import By

class ScraperBot:
    
    def __init__(self):
       
        print("login")
        self.k=0
        edge_options = Options()  # Use EdgeOptions instead of ChromeOptions
        user_agent = 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36 Edg/91.0.864.59'
        
        edge_options.add_argument(f'user-agent={user_agent}')
        #edge_options.add_argument("--headless=new")
        
        self.bot = webdriver.Edge( options=edge_options)  # Specify the path to msedgedriver

    def getLinks(self):
        bot=self.bot
        
        print("login started")
        
        bot.get('https://www.wonder.legal/in/')
        links={}
        page_source = bot.page_source
        soup = BeautifulSoup(page_source,'html.parser')
        a=soup.find_all('div',{'class':'col col50'})
        for b in a:
            c=b.find_all('a')
            for d in c:
                links[d.text]=d['href']
        
        return links
    
    def getData(self, links):
        bot = self.bot
        wait = WebDriverWait(bot, 10) 
        data = {}
        #links=dict(list(links.items())[:2])

        for k, v in links.items():
            bot.get(v)

            # Wait until texte_a_afficher is present (main content loaded)
            wait.until(EC.presence_of_element_located((By.ID, 'texte_a_afficher')))

            soup = BeautifulSoup(bot.page_source, 'html.parser')
            dat = soup.find('div', {'id': 'texte_a_afficher'}).text
            data[k] = [dat]

            questions = []

            while True:
                
                soup = BeautifulSoup(bot.page_source, 'html.parser')
                ques = soup.find_all('label', {'class': 'label_question'})
                for q in ques:
                    questions.append(q.text)

                try:
                    
                    button = wait.until(
                        EC.element_to_be_clickable((By.NAME, 'suivant'))
                    )
                    button.click()
                    
                    wait.until(EC.staleness_of(button))
                    
                except:
                    
                    break
            quest = "\n".join(f"{i+1}. {q}" for i, q in enumerate(questions)) + "\n"
            data[k].append(quest)

        return data

    
    def convert_url(self,url):
        return url.replace("/modele/", "/creation-modele/")
    
    def close(self):
        print("closed")
        self.bot.quit()

st.title("Wonder.Legal Scraper")

if st.button("Start Scraping"):
    start_time=time.time()
    with st.spinner("Scraping in progress..."):
        bot = ScraperBot()
        links = bot.getLinks()
        word_links={}
        for k,v in links.items():
            word_links[k]=bot.convert_url(v)

        st.success(f"Found {len(word_links)} links.")
        
        data = bot.getData(word_links)
        bot.close()

        df = pd.DataFrame([(k, v[0], v[1]) for k, v in data.items()], columns=['Title', 'Description', 'Questions'])
        st.write(df)
        df.to_excel('wonder_legal_data.xlsx', index=False)
        output_folder = "word_files"
        os.makedirs(output_folder, exist_ok=True)

        for index, row in df.iterrows():
            title = row['Title']
            description = row['Description']
            questions = row['Questions']

            # Create a safe folder name
            safe_title = "".join(c if c.isalnum() or c in (' ', '_', '-') else "_" for c in title)
            folder_path = os.path.join(output_folder, safe_title)
            os.makedirs(folder_path, exist_ok=True)

            # Save description.docx
            desc_doc = Document()
            desc_doc.add_paragraph(description)
            desc_path = os.path.join(folder_path, "description.docx")
            desc_doc.save(desc_path)

            # Save questions.docx
            ques_doc = Document()
            if isinstance(questions, list):
                for q in questions:
                    ques_doc.add_paragraph(q)
            else:
                ques_doc.add_paragraph(str(questions))  

            ques_path = os.path.join(folder_path, "questions.docx")
            ques_doc.save(ques_path)
        end_time=time.time()
        elapsed_time=end_time-start_time
        st.success(f"PROCESS COMPLETED in {elapsed_time:.2f} seconds")
